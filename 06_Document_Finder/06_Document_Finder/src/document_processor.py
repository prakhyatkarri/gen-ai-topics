"""Document Processor Module

Handles extraction of text content from PDF, Word, and Markdown files.
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Optional
import re

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import markdown
except ImportError:
    markdown = None


class DocumentProcessor:
    """Process and extract text from various document formats."""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF',
        '.docx': 'Word',
        '.doc': 'Word',
        '.md': 'Markdown',
        '.markdown': 'Markdown'
    }
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        """Initialize DocumentProcessor.
        
        Args:
            logger: Optional logger instance
        """
        self.logger = logger or logging.getLogger(__name__)
        self._validate_dependencies()
    
    def _validate_dependencies(self):
        """Check if required libraries are installed."""
        missing = []
        if PyPDF2 is None:
            missing.append('PyPDF2')
        if Document is None:
            missing.append('python-docx')
        if markdown is None:
            missing.append('markdown')
        
        if missing:
            self.logger.warning(
                f"Some dependencies are missing: {', '.join(missing)}. "
                f"Install them to enable full functionality."
            )
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file format is supported
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_FORMATS
    
    def get_file_type(self, file_path: str) -> Optional[str]:
        """Get the document type from file extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Document type or None if unsupported
        """
        ext = Path(file_path).suffix.lower()
        return self.SUPPORTED_FORMATS.get(ext)
    
    def process_document(self, file_path: str) -> Dict[str, any]:
        """Process a document and extract metadata and content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing document metadata and content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file format: {file_path}")
        
        file_type = self.get_file_type(file_path)
        
        try:
            if file_type == 'PDF':
                content = self._extract_pdf(file_path)
            elif file_type == 'Word':
                content = self._extract_word(file_path)
            elif file_type == 'Markdown':
                content = self._extract_markdown(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_type': file_type,
                'content': content,
                'size': os.path.getsize(file_path),
                'word_count': len(content.split()),
                'char_count': len(content)
            }
        
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
        
        text = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {page_num} from {file_path}: {str(e)}")
                        continue
            
            return "\n".join(text)
        
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
    
    def _extract_word(self, file_path: str) -> str:
        """Extract text from Word document.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Extracted text content
        """
        if Document is None:
            raise ImportError("python-docx is required for Word processing. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            return "\n".join(text)
        
        except Exception as e:
            self.logger.error(f"Error reading Word document {file_path}: {str(e)}")
            raise
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file.
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Option 1: Return raw markdown
            # This preserves the original markdown syntax
            return content
            
            # Option 2: Convert to HTML and strip tags (if markdown library is available)
            # if markdown is not None:
            #     html = markdown.markdown(content)
            #     # Simple HTML tag removal
            #     clean_text = re.sub('<[^<]+?>', '', html)
            #     return clean_text
            # return content
        
        except Exception as e:
            self.logger.error(f"Error reading Markdown file {file_path}: {str(e)}")
            raise
    
    def process_directory(self, directory: str, recursive: bool = True) -> List[Dict[str, any]]:
        """Process all supported documents in a directory.
        
        Args:
            directory: Path to directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of processed document dictionaries
        """
        if not os.path.exists(directory):
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        documents = []
        
        if recursive:
            for root, _, files in os.walk(directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.is_supported(file_path):
                        try:
                            doc = self.process_document(file_path)
                            documents.append(doc)
                            self.logger.info(f"Processed: {file_path}")
                        except Exception as e:
                            self.logger.error(f"Failed to process {file_path}: {str(e)}")
                            continue
        else:
            for file in os.listdir(directory):
                file_path = os.path.join(directory, file)
                if os.path.isfile(file_path) and self.is_supported(file_path):
                    try:
                        doc = self.process_document(file_path)
                        documents.append(doc)
                        self.logger.info(f"Processed: {file_path}")
                    except Exception as e:
                        self.logger.error(f"Failed to process {file_path}: {str(e)}")
                        continue
        
        return documents